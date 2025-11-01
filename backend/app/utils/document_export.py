"""
文档导出工具函数
"""
import os
import re
import uuid
import logging
import tempfile
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path

# 导入所需的库
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logging.warning("WeasyPrint not available, PDF export will be limited")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available, Word export will be limited")

try:
    import markdown2
    MARKDOWN2_AVAILABLE = True
except ImportError:
    MARKDOWN2_AVAILABLE = False
    logging.warning("markdown2 not available, Markdown export will be limited")

try:
    from jinja2 import Environment, BaseLoader
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False
    logging.warning("Jinja2 not available, template rendering will be limited")

from app.models.document import DocumentFormat
from app.schemas.document_export import ExportFormat, ExportOptions

# 设置日志
logger = logging.getLogger(__name__)

# 导出文件存储目录
EXPORT_DIR = os.environ.get("EXPORT_DIR", "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


class DocumentExportError(Exception):
    """文档导出异常"""
    pass


def convert_markdown_to_html(markdown_content: str) -> str:
    """将Markdown转换为HTML"""
    if not MARKDOWN2_AVAILABLE:
        # 简单的Markdown到HTML转换
        html_content = markdown_content
        html_content = re.sub(r'### (.*)', r'<h3>\1</h3>', html_content)
        html_content = re.sub(r'## (.*)', r'<h2>\1</h2>', html_content)
        html_content = re.sub(r'# (.*)', r'<h1>\1</h1>', html_content)
        html_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_content)
        html_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_content)
        html_content = re.sub(r'`([^`]*)`', r'<code>\1</code>', html_content)
        html_content = re.sub(r'\n\n', '</p><p>', html_content)
        html_content = '<p>' + html_content + '</p>'
        html_content = re.sub(r'<p></p>', '', html_content)
        return html_content
    
    # 使用markdown2库进行转换
    extras = ["tables", "fenced-code-blocks", "header-ids", "toc"]
    return markdown2.markdown(markdown_content, extras=extras)


def convert_html_to_markdown(html_content: str) -> str:
    """将HTML转换为Markdown"""
    # 简单的HTML到Markdown转换
    markdown_content = html_content
    markdown_content = re.sub(r'<h1>(.*?)</h1>', r'# \1', markdown_content)
    markdown_content = re.sub(r'<h2>(.*?)</h2>', r'## \1', markdown_content)
    markdown_content = re.sub(r'<h3>(.*?)</h3>', r'### \1', markdown_content)
    markdown_content = re.sub(r'<strong>(.*?)</strong>', r'**\1**', markdown_content)
    markdown_content = re.sub(r'<em>(.*?)</em>', r'*\1*', markdown_content)
    markdown_content = re.sub(r'<code>(.*?)</code>', r'`\1`', markdown_content)
    markdown_content = re.sub(r'<p>(.*?)</p>', r'\1\n\n', markdown_content)
    markdown_content = re.sub(r'\n\n+', '\n\n', markdown_content)
    return markdown_content.strip()


def convert_to_plain_text(content: str, source_format: DocumentFormat) -> str:
    """将内容转换为纯文本"""
    if source_format == DocumentFormat.HTML:
        # 移除HTML标签
        content = re.sub(r'<[^>]+>', '', content)
    elif source_format == DocumentFormat.MARKDOWN:
        # 移除Markdown标记
        content = re.sub(r'[#*`]', '', content)
    return content


def generate_pdf(
    content: str,
    title: str,
    options: ExportOptions,
    source_format: DocumentFormat = DocumentFormat.HTML
) -> Tuple[str, int]:
    """生成PDF文件"""
    if not WEASYPRINT_AVAILABLE:
        raise DocumentExportError("WeasyPrint not available for PDF generation")
    
    try:
        # 根据源格式转换内容
        if source_format == DocumentFormat.MARKDOWN:
            html_content = convert_markdown_to_html(content)
        elif source_format == DocumentFormat.PLAIN_TEXT:
            html_content = f"<p>{content.replace(chr(10), '<br>')}</p>"
        else:
            html_content = content
        
        # 创建HTML模板
        if JINJA2_AVAILABLE:
            template_str = """
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{{ title }}</title>
                <style>
                    @page {
                        size: {{ pdf_page_size }};
                        margin: {{ pdf_margin_top }}cm {{ pdf_margin_right }}cm {{ pdf_margin_bottom }}cm {{ pdf_margin_left }}cm;
                    }
                    body {
                        font-family: Arial, sans-serif;
                        font-size: {{ pdf_font_size }}pt;
                        line-height: {{ pdf_line_height }};
                        color: #333;
                    }
                    h1, h2, h3, h4, h5, h6 {
                        color: #2c3e50;
                        margin-top: 1.5em;
                        margin-bottom: 0.5em;
                    }
                    h1 { font-size: 2em; }
                    h2 { font-size: 1.5em; }
                    h3 { font-size: 1.2em; }
                    p {
                        margin-bottom: 1em;
                        text-align: justify;
                    }
                    code {
                        background-color: #f8f9fa;
                        padding: 2px 4px;
                        border-radius: 3px;
                        font-family: monospace;
                    }
                    pre {
                        background-color: #f8f9fa;
                        padding: 1em;
                        border-radius: 5px;
                        overflow-x: auto;
                    }
                    blockquote {
                        border-left: 4px solid #ddd;
                        padding-left: 1em;
                        margin-left: 0;
                        color: #666;
                    }
                    table {
                        border-collapse: collapse;
                        width: 100%;
                        margin-bottom: 1em;
                    }
                    th, td {
                        border: 1px solid #ddd;
                        padding: 8px;
                        text-align: left;
                    }
                    th {
                        background-color: #f2f2f2;
                    }
                    .header, .footer {
                        font-size: 10pt;
                        color: #666;
                    }
                    .header {
                        text-align: center;
                        margin-bottom: 2em;
                    }
                    .footer {
                        text-align: center;
                        margin-top: 2em;
                    }
                    .watermark {
                        position: fixed;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%) rotate(-45deg);
                        font-size: 72pt;
                        color: rgba(200, 200, 200, {{ watermark_opacity }});
                        z-index: -1;
                    }
                </style>
            </head>
            <body>
                {% if include_watermark and watermark_text %}
                <div class="watermark">{{ watermark_text }}</div>
                {% endif %}
                
                {% if include_header %}
                <div class="header">
                    {{ header or title }}
                </div>
                {% endif %}
                
                <h1>{{ title }}</h1>
                
                {{ html_content|safe }}
                
                {% if include_footer %}
                <div class="footer">
                    {{ footer or "Generated on " + current_date }}
                </div>
                {% endif %}
            </body>
            </html>
            """
            
            env = Environment(loader=BaseLoader())
            template = env.from_string(template_str)
            
            # 渲染模板
            html_content = template.render(
                title=title,
                html_content=html_content,
                current_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                pdf_page_size=options.pdf_page_size,
                pdf_margin_top=options.pdf_margin_top,
                pdf_margin_bottom=options.pdf_margin_bottom,
                pdf_margin_left=options.pdf_margin_left,
                pdf_margin_right=options.pdf_margin_right,
                pdf_font_size=options.pdf_font_size,
                pdf_line_height=options.pdf_line_height,
                include_header=options.include_header,
                include_footer=options.include_footer,
                header=options.watermark_text if options.include_watermark else None,
                footer=options.watermark_text if options.include_watermark else None,
                include_watermark=options.include_watermark,
                watermark_text=options.watermark_text,
                watermark_opacity=options.watermark_opacity
            )
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_name = f"{title}_{file_id}.pdf"
        file_path = os.path.join(EXPORT_DIR, file_name)
        
        # 生成PDF
        html_doc = HTML(string=html_content)
        html_doc.write_pdf(file_path)
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        logger.info(f"PDF generated successfully: {file_path}")
        return file_path, file_size
        
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise DocumentExportError(f"Failed to generate PDF: {str(e)}")


def generate_word(
    content: str,
    title: str,
    options: ExportOptions,
    source_format: DocumentFormat = DocumentFormat.HTML
) -> Tuple[str, int]:
    """生成Word文档"""
    if not PYTHON_DOCX_AVAILABLE:
        raise DocumentExportError("python-docx not available for Word generation")
    
    try:
        # 创建Word文档
        doc = Document()
        
        # 设置页面样式
        section = doc.sections[0]
        section.page_height = Inches(11.7) if options.word_page_size == "A4" else Inches(11.0)
        section.page_width = Inches(8.3) if options.word_page_size == "A4" else Inches(8.5)
        section.left_margin = Inches(options.word_margin_left / 2.54)
        section.right_margin = Inches(options.word_margin_right / 2.54)
        section.top_margin = Inches(options.word_margin_top / 2.54)
        section.bottom_margin = Inches(options.word_margin_bottom / 2.54)
        
        # 添加标题
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加页眉
        if options.include_header:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = options.watermark_text if options.include_watermark else title
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加页脚
        if options.include_footer:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加水印
        if options.include_watermark and options.watermark_text:
            # 创建水印
            watermark = doc.sections[0].header
            watermark_para = watermark.paragraphs[0] if watermark.paragraphs else watermark.add_paragraph()
            watermark_run = watermark_para.add_run(options.watermark_text)
            watermark_run.font.size = Pt(72)
            watermark_run.font.color.rgb = None  # 设置为透明或浅色
            watermark_run.bold = True
        
        # 根据源格式处理内容
        if source_format == DocumentFormat.MARKDOWN:
            # 简单处理Markdown内容
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip() == '':
                    doc.add_paragraph()  # 空行
                else:
                    # 处理粗体和斜体
                    line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # 移除粗体标记
                    line = re.sub(r'\*(.*?)\*', r'\1', line)  # 移除斜体标记
                    line = re.sub(r'`([^`]*)`', r'\1', line)  # 移除代码标记
                    doc.add_paragraph(line)
        elif source_format == DocumentFormat.HTML:
            # 简单处理HTML内容
            # 移除HTML标签，保留基本格式
            clean_content = re.sub(r'<[^>]+>', '', content)
            paragraphs = clean_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        else:
            # 纯文本内容
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_name = f"{title}_{file_id}.docx"
        file_path = os.path.join(EXPORT_DIR, file_name)
        
        # 保存文档
        doc.save(file_path)
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        logger.info(f"Word document generated successfully: {file_path}")
        return file_path, file_size
        
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        raise DocumentExportError(f"Failed to generate Word document: {str(e)}")


def generate_html(
    content: str,
    title: str,
    options: ExportOptions,
    source_format: DocumentFormat = DocumentFormat.HTML
) -> Tuple[str, int]:
    """生成HTML文件"""
    try:
        # 根据源格式转换内容
        if source_format == DocumentFormat.MARKDOWN:
            html_content = convert_markdown_to_html(content)
        elif source_format == DocumentFormat.PLAIN_TEXT:
            html_content = f"<p>{content.replace(chr(10), '<br>')}</p>"
        else:
            html_content = content
        
        # 创建HTML模板
        if JINJA2_AVAILABLE:
            template_str = """
            <!DOCTYPE html>
            <html lang="zh-CN">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{{ title }}</title>
                <style>
                    body {
                        font-family: Arial, sans-serif;
                        line-height: {{ word_line_height }};
                        color: #333;
                        max-width: 800px;
                        margin: 0 auto;
                        padding: 20px;
                        {% if html_responsive %}
                        @media (max-width: 600px) {
                            body {
                                padding: 10px;
                            }
                        }
                        {% endif %}
                    }
                    h1, h2, h3, h4, h5, h6 {
                        color: #2c3e50;
                        margin-top: 1.5em;
                        margin-bottom: 0.5em;
                    }
                    h1 { font-size: 2em; }
                    h2 { font-size: 1.5em; }
                    h3 { font-size: 1.2em; }
                    p {
                        margin-bottom: 1em;
                        text-align: justify;
                    }
                    code {
                        background-color: #f8f9fa;
                        padding: 2px 4px;
                        border-radius: 3px;
                        font-family: monospace;
                    }
                    pre {
                        background-color: #f8f9fa;
                        padding: 1em;
                        border-radius: 5px;
                        overflow-x: auto;
                    }
                    blockquote {
                        border-left: 4px solid #ddd;
                        padding-left: 1em;
                        margin-left: 0;
                        color: #666;
                    }
                    table {
                        border-collapse: collapse;
                        width: 100%;
                        margin-bottom: 1em;
                    }
                    th, td {
                        border: 1px solid #ddd;
                        padding: 8px;
                        text-align: left;
                    }
                    th {
                        background-color: #f2f2f2;
                    }
                    .header, .footer {
                        font-size: 10pt;
                        color: #666;
                        text-align: center;
                        margin: 1em 0;
                    }
                    .watermark {
                        position: fixed;
                        top: 50%;
                        left: 50%;
                        transform: translate(-50%, -50%) rotate(-45deg);
                        font-size: 72pt;
                        color: rgba(200, 200, 200, {{ watermark_opacity }});
                        z-index: -1;
                        pointer-events: none;
                    }
                    .toc {
                        background-color: #f8f9fa;
                        border: 1px solid #ddd;
                        padding: 1em;
                        margin-bottom: 2em;
                        border-radius: 5px;
                    }
                    .toc h2 {
                        margin-top: 0;
                        font-size: 1.2em;
                    }
                    .toc ul {
                        padding-left: 20px;
                    }
                    .toc li {
                        margin-bottom: 0.5em;
                    }
                </style>
            </head>
            <body>
                {% if include_watermark and watermark_text %}
                <div class="watermark">{{ watermark_text }}</div>
                {% endif %}
                
                {% if include_header %}
                <div class="header">
                    {{ header or title }}
                </div>
                {% endif %}
                
                <h1>{{ title }}</h1>
                
                {% if html_include_toc %}
                <div class="toc">
                    <h2>目录</h2>
                    <ul>
                        <li><a href="#section-1">第一节</a></li>
                        <li><a href="#section-2">第二节</a></li>
                        <!-- 这里应该根据实际内容生成目录 -->
                    </ul>
                </div>
                {% endif %}
                
                {{ html_content|safe }}
                
                {% if include_footer %}
                <div class="footer">
                    {{ footer or "Generated on " + current_date }}
                </div>
                {% endif %}
            </body>
            </html>
            """
            
            env = Environment(loader=BaseLoader())
            template = env.from_string(template_str)
            
            # 渲染模板
            html_content = template.render(
                title=title,
                html_content=html_content,
                current_date=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                word_line_height=options.word_line_height,
                html_responsive=options.html_responsive,
                include_header=options.include_header,
                include_footer=options.include_footer,
                header=options.watermark_text if options.include_watermark else None,
                footer=options.watermark_text if options.include_watermark else None,
                include_watermark=options.include_watermark,
                watermark_text=options.watermark_text,
                watermark_opacity=options.watermark_opacity,
                html_include_toc=options.html_include_toc
            )
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_name = f"{title}_{file_id}.html"
        file_path = os.path.join(EXPORT_DIR, file_name)
        
        # 保存HTML文件
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        logger.info(f"HTML file generated successfully: {file_path}")
        return file_path, file_size
        
    except Exception as e:
        logger.error(f"Error generating HTML file: {str(e)}")
        raise DocumentExportError(f"Failed to generate HTML file: {str(e)}")


def generate_markdown(
    content: str,
    title: str,
    options: ExportOptions,
    source_format: DocumentFormat = DocumentFormat.HTML
) -> Tuple[str, int]:
    """生成Markdown文件"""
    try:
        # 根据源格式转换内容
        if source_format == DocumentFormat.HTML:
            markdown_content = convert_html_to_markdown(content)
        elif source_format == DocumentFormat.PLAIN_TEXT:
            markdown_content = content
        else:
            markdown_content = content
        
        # 添加标题
        if not markdown_content.startswith(f"# {title}"):
            markdown_content = f"# {title}\n\n{markdown_content}"
        
        # 添加元数据
        metadata = []
        metadata.append(f"---")
        metadata.append(f"title: {title}")
        metadata.append(f"generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if options.include_watermark and options.watermark_text:
            metadata.append(f"watermark: {options.watermark_text}")
        metadata.append(f"---")
        metadata.append("")
        
        # 组合内容
        final_content = "\n".join(metadata) + markdown_content
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_name = f"{title}_{file_id}.md"
        file_path = os.path.join(EXPORT_DIR, file_name)
        
        # 保存Markdown文件
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        # 获取文件大小
        file_size = os.path.getsize(file_path)
        
        logger.info(f"Markdown file generated successfully: {file_path}")
        return file_path, file_size
        
    except Exception as e:
        logger.error(f"Error generating Markdown file: {str(e)}")
        raise DocumentExportError(f"Failed to generate Markdown file: {str(e)}")


def export_document(
    content: str,
    title: str,
    export_format: ExportFormat,
    options: ExportOptions,
    source_format: DocumentFormat = DocumentFormat.HTML
) -> Tuple[str, int]:
    """
    导出文档
    
    Args:
        content: 文档内容
        title: 文档标题
        export_format: 导出格式
        options: 导出选项
        source_format: 源格式
    
    Returns:
        Tuple[str, int]: (文件路径, 文件大小)
    
    Raises:
        DocumentExportError: 导出失败时抛出
    """
    try:
        if export_format == ExportFormat.PDF:
            return generate_pdf(content, title, options, source_format)
        elif export_format == ExportFormat.WORD:
            return generate_word(content, title, options, source_format)
        elif export_format == ExportFormat.HTML:
            return generate_html(content, title, options, source_format)
        elif export_format == ExportFormat.MARKDOWN:
            return generate_markdown(content, title, options, source_format)
        else:
            raise DocumentExportError(f"Unsupported export format: {export_format}")
    except Exception as e:
        logger.error(f"Error exporting document: {str(e)}")
        raise DocumentExportError(f"Failed to export document: {str(e)}")


def cleanup_export_file(file_path: str) -> bool:
    """
    清理导出文件
    
    Args:
        file_path: 文件路径
    
    Returns:
        bool: 清理是否成功
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Export file cleaned up: {file_path}")
            return True
        return False
    except Exception as e:
        logger.error(f"Error cleaning up export file: {str(e)}")
        return False


def get_export_file_url(file_path: str) -> str:
    """
    获取导出文件的URL
    
    Args:
        file_path: 文件路径
    
    Returns:
        str: 文件URL
    """
    # 获取相对路径
    abs_path = os.path.abspath(file_path)
    export_dir_abs = os.path.abspath(EXPORT_DIR)
    
    if abs_path.startswith(export_dir_abs):
        rel_path = os.path.relpath(abs_path, export_dir_abs)
        return f"/exports/{rel_path}"
    
    # 如果不在导出目录中，返回绝对路径
    return file_path


def validate_export_options(options: ExportOptions) -> bool:
    """
    验证导出选项
    
    Args:
        options: 导出选项
    
    Returns:
        bool: 验证是否通过
    """
    try:
        # 验证PDF选项
        if options.pdf_page_size not in ["A4", "A3", "Letter", "Legal"]:
            return False
        
        if options.pdf_orientation not in ["portrait", "landscape"]:
            return False
        
        # 验证Word选项
        if options.word_page_size not in ["A4", "A3", "Letter", "Legal"]:
            return False
        
        if options.word_orientation not in ["portrait", "landscape"]:
            return False
        
        # 验证边距
        if any(m < 0 for m in [
            options.pdf_margin_top, options.pdf_margin_bottom,
            options.pdf_margin_left, options.pdf_margin_right,
            options.word_margin_top, options.word_margin_bottom,
            options.word_margin_left, options.word_margin_right
        ]):
            return False
        
        # 验证字体大小
        if options.pdf_font_size <= 0 or options.word_font_size <= 0:
            return False
        
        # 验证行高
        if options.pdf_line_height <= 0 or options.word_line_height <= 0:
            return False
        
        # 验证水印透明度
        if not (0.0 <= options.watermark_opacity <= 1.0):
            return False
        
        return True
    except Exception:
        return False