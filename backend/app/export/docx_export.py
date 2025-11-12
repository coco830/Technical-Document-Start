"""
Word 文档导出模块
使用 python-docx 生成 Word 文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
from datetime import datetime
from pathlib import Path
import re
from typing import Optional, Dict, Any


class DocxExporter:
    """Word 文档导出器"""

    def __init__(self, output_dir: str = "exports/docx"):
        """
        初始化 Word 导出器

        Args:
            output_dir: 输出目录路径
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _setup_styles(self, doc: Document):
        """
        设置文档样式

        Args:
            doc: Word 文档对象
        """
        try:
            # 标题样式
            if 'CustomTitle' not in doc.styles:
                title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
                title_font = title_style.font
                title_font.name = 'Arial'
                title_font.size = Pt(24)
                title_font.bold = True
                title_font.color.rgb = RGBColor(26, 26, 26)
                title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_style.paragraph_format.space_after = Pt(18)

            # 一级标题样式
            if 'CustomHeading1' not in doc.styles:
                h1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                h1_font = h1_style.font
                h1_font.name = 'Arial'
                h1_font.size = Pt(18)
                h1_font.bold = True
                h1_font.color.rgb = RGBColor(44, 62, 80)
                h1_style.paragraph_format.space_before = Pt(12)
                h1_style.paragraph_format.space_after = Pt(6)

            # 二级标题样式
            if 'CustomHeading2' not in doc.styles:
                h2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
                h2_font = h2_style.font
                h2_font.name = 'Arial'
                h2_font.size = Pt(14)
                h2_font.bold = True
                h2_font.color.rgb = RGBColor(52, 73, 94)
                h2_style.paragraph_format.space_before = Pt(10)
                h2_style.paragraph_format.space_after = Pt(6)

            # 正文样式
            if 'CustomBody' not in doc.styles:
                body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
                body_font = body_style.font
                body_font.name = 'Arial'
                body_font.size = Pt(11)
                body_style.paragraph_format.line_spacing = 1.5
                body_style.paragraph_format.space_after = Pt(8)
                body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # 元数据样式
            if 'CustomMetadata' not in doc.styles:
                meta_style = doc.styles.add_style('CustomMetadata', WD_STYLE_TYPE.PARAGRAPH)
                meta_font = meta_style.font
                meta_font.name = 'Arial'
                meta_font.size = Pt(9)
                meta_font.color.rgb = RGBColor(127, 140, 141)
                meta_style.paragraph_format.space_after = Pt(6)

        except Exception as e:
            print(f"样式设置警告: {e}")

    def _html_to_text(self, html_content: str) -> str:
        """
        将 HTML 转换为纯文本

        Args:
            html_content: HTML 内容

        Returns:
            纯文本内容
        """
        if not html_content:
            return ""

        try:
            soup = BeautifulSoup(html_content, 'html.parser')

            # 移除 script 和 style 标签
            for script in soup(["script", "style"]):
                script.decompose()

            # 获取文本
            text = soup.get_text(separator='\n')

            # 清理多余的空行
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return '\n\n'.join(lines)
        except Exception as e:
            print(f"HTML 解析失败: {e}")
            return html_content

    def _generate_filename(self, title: str, timestamp: bool = True) -> str:
        """
        生成文件名

        Args:
            title: 文档标题
            timestamp: 是否添加时间戳

        Returns:
            文件名
        """
        # 清理标题中的特殊字符
        clean_title = re.sub(r'[<>:"/\\|?*]', '_', title)
        clean_title = clean_title[:50]  # 限制长度

        if timestamp:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{clean_title}_{ts}.docx"
        else:
            filename = f"{clean_title}.docx"

        return filename

    def _add_content_to_doc(self, doc: Document, content: str, content_type: str = "text"):
        """
        添加内容到文档

        Args:
            doc: Word 文档对象
            content: 内容文本
            content_type: 内容类型
        """
        # 处理内容
        if content_type == "html":
            text_content = self._html_to_text(content)
        else:
            text_content = content

        # 分段处理
        paragraphs = text_content.split('\n\n')

        for para_text in paragraphs:
            if not para_text.strip():
                continue

            # 检测标题（简单规则）
            if para_text.strip().startswith('#'):
                # Markdown 风格标题
                if para_text.startswith('###'):
                    para = doc.add_paragraph(para_text.replace('#', '').strip())
                    para.style = 'CustomHeading2'
                elif para_text.startswith('##'):
                    para = doc.add_paragraph(para_text.replace('#', '').strip())
                    para.style = 'CustomHeading1'
                else:
                    para = doc.add_paragraph(para_text.replace('#', '').strip())
                    para.style = 'CustomTitle'
            else:
                # 普通段落
                para = doc.add_paragraph(para_text.strip())
                para.style = 'CustomBody'

    def export(
        self,
        title: str,
        content: str,
        content_type: str = "html",
        metadata: Optional[Dict[str, Any]] = None,
        custom_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        导出文档为 Word

        Args:
            title: 文档标题
            content: 文档内容
            content_type: 内容类型 (html, markdown, text)
            metadata: 文档元数据
            custom_filename: 自定义文件名

        Returns:
            导出结果字典
        """
        try:
            # 创建文档
            doc = Document()

            # 设置样式
            self._setup_styles(doc)

            # 设置页边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # 添加标题
            title_para = doc.add_paragraph(title)
            title_para.style = 'CustomTitle'

            # 添加空行
            doc.add_paragraph()

            # 添加元数据
            if metadata:
                metadata_parts = []
                if metadata.get('author'):
                    metadata_parts.append(f"作者: {metadata['author']}")
                if metadata.get('created_at'):
                    metadata_parts.append(f"创建时间: {metadata['created_at']}")
                if metadata.get('version'):
                    metadata_parts.append(f"版本: {metadata['version']}")

                if metadata_parts:
                    meta_para = doc.add_paragraph(' | '.join(metadata_parts))
                    meta_para.style = 'CustomMetadata'

            # 添加导出时间戳
            export_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            export_para = doc.add_paragraph(f"导出时间: {export_time}")
            export_para.style = 'CustomMetadata'

            # 添加分隔线
            doc.add_paragraph('_' * 60)
            doc.add_paragraph()

            # 添加内容
            self._add_content_to_doc(doc, content, content_type)

            # 生成文件名
            filename = custom_filename if custom_filename else self._generate_filename(title)
            filepath = self.output_dir / filename

            # 保存文档
            doc.save(str(filepath))

            # 验证文件是否存在
            if not filepath.exists():
                raise Exception("Word 文件生成失败")

            file_size = filepath.stat().st_size

            return {
                'success': True,
                'filename': filename,
                'filepath': str(filepath),
                'file_size': file_size,
                'export_time': export_time,
                'message': 'Word 文档导出成功'
            }

        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f'Word 文档导出失败: {str(e)}'
            }

    def export_batch(
        self,
        documents: list,
        output_filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        批量导出文档为单个 Word 文档

        Args:
            documents: 文档列表，每个文档包含 title, content, content_type
            output_filename: 输出文件名

        Returns:
            导出结果字典
        """
        try:
            # 创建文档
            doc = Document()

            # 设置样式
            self._setup_styles(doc)

            # 设置页边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # 添加封面
            cover_title = doc.add_paragraph("批量文档导出")
            cover_title.style = 'CustomTitle'

            doc.add_paragraph()

            export_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            meta1 = doc.add_paragraph(f"导出时间: {export_time}")
            meta1.style = 'CustomMetadata'

            meta2 = doc.add_paragraph(f"文档数量: {len(documents)}")
            meta2.style = 'CustomMetadata'

            # 分页
            doc.add_page_break()

            # 添加每个文档
            for idx, doc_data in enumerate(documents, 1):
                # 文档标题
                doc_title = doc.add_paragraph(f"{idx}. {doc_data['title']}")
                doc_title.style = 'CustomTitle'

                doc.add_paragraph()

                # 添加内容
                content_type = doc_data.get('content_type', 'html')
                content = doc_data.get('content', '')
                self._add_content_to_doc(doc, content, content_type)

                # 文档分隔（最后一个文档不分页）
                if idx < len(documents):
                    doc.add_page_break()

            # 生成文件名
            if not output_filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"batch_export_{timestamp}.docx"

            filepath = self.output_dir / output_filename

            # 保存文档
            doc.save(str(filepath))

            # 验证文件
            if not filepath.exists():
                raise Exception("批量 Word 文件生成失败")

            file_size = filepath.stat().st_size

            return {
                'success': True,
                'filename': output_filename,
                'filepath': str(filepath),
                'file_size': file_size,
                'document_count': len(documents),
                'export_time': export_time,
                'message': '批量 Word 文档导出成功'
            }

        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f'批量 Word 文档导出失败: {str(e)}'
            }
